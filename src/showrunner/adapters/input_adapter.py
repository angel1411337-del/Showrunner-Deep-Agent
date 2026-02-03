"""Input adapters for loading documents from files and folders.

Provides unified interface for loading single files or folders of chapter files,
normalizing both to the same DocumentUnit model for downstream processing.
"""

import warnings
from pathlib import Path
from typing import Protocol, cast

from showrunner.contracts import DocumentUnit


class InputAdapter(Protocol):
    """Protocol for input adapters that load documents from various sources."""

    def load(self, source: Path) -> list[DocumentUnit]:
        """Load and normalize input to DocumentUnits.

        Args:
            source: Path to the input source (file or folder).

        Returns:
            List of DocumentUnit objects normalized from the source.
        """
        ...

    def load_files(self, files: list[Path]) -> list[DocumentUnit]:
        """Load a specific set of files as DocumentUnits.

        Args:
            files: List of files to load.

        Returns:
            List of DocumentUnit objects.
        """
        ...


SUPPORTED_EXTENSIONS: set[str] = {".txt", ".md", ".markdown", ".docx", ".pdf"}


class _DocxParagraph(Protocol):
    text: str


class _DocxDocument(Protocol):
    paragraphs: list[_DocxParagraph]


class _PdfPage(Protocol):
    def extract_text(self) -> str | None: ...


class _PdfReader(Protocol):
    pages: list[_PdfPage]


def _load_text_file(source: Path) -> str:
    return source.read_text(encoding="utf-8")


def _load_docx_file(source: Path) -> str:
    try:
        import docx  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required to read .docx files") from exc

    document = cast(
        "_DocxDocument",
        docx.Document(source),  # type: ignore[reportUnknownMemberType]
    )
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    return "\n".join(paragraphs)


def _load_pdf_file(source: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except ImportError as exc:
            raise RuntimeError("pypdf is required to read .pdf files") from exc

    reader = cast("_PdfReader", PdfReader(str(source)))
    parts: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        parts.append(f"=== Page {page_index} ===\n{text}")
    return "\n\n".join(parts)


def _load_file_text(source: Path) -> str:
    ext = source.suffix.lower()
    if ext in {".txt", ".md", ".markdown"}:
        return _load_text_file(source)
    if ext == ".docx":
        return _load_docx_file(source)
    if ext == ".pdf":
        return _load_pdf_file(source)
    raise ValueError(f"Unsupported file type: {source.suffix}")


def parse_filename_metadata(filename: str) -> tuple[str | None, str | None]:
    """Extract book and chapter labels from a filename.

    Parsing rules:
    - book1_chapter01_Bran.txt -> book_label="book1", chapter_label="chapter01_Bran"
    - chapter_01.txt -> book_label=None, chapter_label="chapter_01"
    - prologue.txt -> book_label=None, chapter_label="prologue"

    Args:
        filename: The filename (with or without extension) to parse.

    Returns:
        Tuple of (book_label, chapter_label). book_label may be None.
    """
    # Remove extension if present
    stem = Path(filename).stem if "." in filename else filename

    # Check for book_chapter pattern: book<N>_<rest>
    if stem.startswith("book") and "_" in stem:
        parts = stem.split("_", 1)
        book_part = parts[0]
        # Verify book_part matches pattern book<digits>
        if book_part[4:].isdigit() or len(book_part) > 4:
            book_label = book_part
            chapter_label = parts[1] if len(parts) > 1 else None
            return book_label, chapter_label

    # No book label - entire stem is chapter label
    return None, stem


class FileInputAdapter:
    """Adapter for loading a single file input.

    Loads a single file and normalizes it to a DocumentUnit.
    """

    def load(self, source: Path) -> list[DocumentUnit]:
        """Load a single file and return as a list with one DocumentUnit.

        Args:
            source: Path to the text file to load.

        Returns:
            List containing a single DocumentUnit.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the path is a directory, not a file.
        """
        if not source.exists():
            raise FileNotFoundError(f"File not found: {source}")

        if source.is_dir():
            raise ValueError(f"Path is not a file: {source}")

        raw_text = _load_file_text(source)
        book_label, chapter_label = parse_filename_metadata(source.name)

        doc = DocumentUnit(
            source_id=source.stem,
            source_path=str(source),
            order_hint=0,
            raw_text=raw_text,
            book_label=book_label,
            chapter_label=chapter_label,
        )

        return [doc]

    def load_files(self, files: list[Path]) -> list[DocumentUnit]:
        """Load a list of files as DocumentUnits."""
        documents: list[DocumentUnit] = []
        for path in files:
            documents.extend(self.load(path))
        return documents


class FolderInputAdapter:
    """Adapter for loading a folder of chapter files.

    Loads supported files from a folder, sorted alphabetically by filename,
    and normalizes each to a DocumentUnit with deterministic ordering.
    """

    def load(self, source: Path) -> list[DocumentUnit]:
        """Load all .txt files from a folder as DocumentUnits.

        Args:
            source: Path to the folder containing chapter files.

        Returns:
            List of DocumentUnit objects, one per .txt file,
            sorted alphabetically by filename with order_hint set accordingly.

        Raises:
            FileNotFoundError: If the folder does not exist.
            ValueError: If the path is a file, not a directory.
        """
        if not source.exists():
            raise FileNotFoundError(f"Folder not found: {source}")

        if source.is_file():
            raise ValueError(f"Path is not a directory: {source}")

        all_files: list[Path] = [f for f in source.iterdir() if f.is_file()]
        supported_files: list[Path] = []
        for file_path in all_files:
            if file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                supported_files.append(file_path)
            else:
                warnings.warn(
                    f"Unsupported file skipped: {file_path.name}",
                    UserWarning,
                    stacklevel=2,
                )
        supported_files = sorted(supported_files, key=lambda p: p.name)

        documents: list[DocumentUnit] = []

        order_hint = 0
        for file_path in supported_files:
            try:
                raw_text = _load_file_text(file_path)
            except RuntimeError as exc:
                warnings.warn(
                    f"Skipping {file_path.name}: {exc}",
                    UserWarning,
                    stacklevel=2,
                )
                continue
            book_label, chapter_label = parse_filename_metadata(file_path.name)

            doc = DocumentUnit(
                source_id=file_path.stem,
                source_path=str(file_path),
                order_hint=order_hint,
                raw_text=raw_text,
                book_label=book_label,
                chapter_label=chapter_label,
            )
            documents.append(doc)
            order_hint += 1

        return documents

    def load_files(self, files: list[Path]) -> list[DocumentUnit]:
        """Load a list of files as DocumentUnits (used for incremental runs)."""
        supported_files: list[Path] = []
        for file_path in files:
            if file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                supported_files.append(file_path)
            else:
                warnings.warn(
                    f"Unsupported file skipped: {file_path.name}",
                    UserWarning,
                    stacklevel=2,
                )
        supported_files = sorted(supported_files, key=lambda p: p.name)

        documents: list[DocumentUnit] = []
        order_hint = 0
        for file_path in supported_files:
            try:
                raw_text = _load_file_text(file_path)
            except RuntimeError as exc:
                warnings.warn(
                    f"Skipping {file_path.name}: {exc}",
                    UserWarning,
                    stacklevel=2,
                )
                continue
            book_label, chapter_label = parse_filename_metadata(file_path.name)
            doc = DocumentUnit(
                source_id=file_path.stem,
                source_path=str(file_path),
                order_hint=order_hint,
                raw_text=raw_text,
                book_label=book_label,
                chapter_label=chapter_label,
            )
            documents.append(doc)
            order_hint += 1

        return documents


def create_adapter(source: Path) -> InputAdapter:
    """Factory to create the appropriate adapter based on source type.

    Args:
        source: Path to the input source.

    Returns:
        FileInputAdapter if source is a file, FolderInputAdapter if directory.

    Raises:
        FileNotFoundError: If the source path does not exist.
    """
    if not source.exists():
        raise FileNotFoundError(f"Source path not found: {source}")

    if source.is_file():
        return FileInputAdapter()
    else:
        return FolderInputAdapter()
